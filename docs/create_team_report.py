from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('PRACTICE SCHOOL – I (FINAL REPORT / REPORT III)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Header Info
doc.add_paragraph('Project: AI-Based Student Mental Health Monitoring and Support System (AIRA)', style='Heading 2')
doc.add_paragraph('Institution: JK Lakshmipat University', style='Heading 3')
doc.add_paragraph('Supervisor: Dr. Sonali Vyas', style='Heading 3')
doc.add_paragraph('Submitted To Internal Supervisor: Mr. Santosh Kumar Verma', style='Heading 3')
doc.add_paragraph('Submitted By: Diksha Shekhawat (2024BTECH156) & Anand Singh Rathore (2024BTECH158)', style='Heading 3')

doc.add_heading('1. INTRODUCTION', level=2)
doc.add_paragraph('This report serves as a direct sequel to the Mid-Term Milestone Report and presents the final phase progress and current condition of the internship project titled AIRA. While the milestone report covered the initial research, problem definition, and early prototype stage, this report details the successful implementation of the core AI architecture, backend services, and the transition from a prototype to a fully functional application.')

doc.add_heading('2. PROGRESS FROM MILESTONE STAGE', level=2)
doc.add_paragraph('At the milestone stage, the system was an early prototype with partial UI implementation and pending AI integration. Since then, the team has successfully completed the objectives outlined in the "Future Work" of the previous report:')
doc.add_paragraph('AI Emotion Detection Models: Successfully trained and integrated both the Text Analysis (NLP) model and the Behavioral ML regression model.', style='List Bullet')
doc.add_paragraph('Backend Integration: The Flask backend API is now fully developed, featuring JWT authentication, validation middleware, and complex routing.', style='List Bullet')
doc.add_paragraph('Database Connectivity: MongoDB collections have been structured and connected to handle real-time user data, mental health reports, and mood logs.', style='List Bullet')
doc.add_paragraph('Visualization Dashboard: The UI/UX has been polished with responsive Chart.js data visualizations and the 30-Day Mood Stability Heatmap.', style='List Bullet')

doc.add_heading('3. WORK COMPLETED IN THE FINAL PHASE', level=2)
doc.add_heading('3.1 Machine Learning & Diagnostic Systems', level=3)
doc.add_paragraph('Integrated a DistilBERT-based NLP pipeline to parse free-text journal logs and accurately classify underlying emotional sentiment (joy, melancholy, anxiety).', style='List Bullet')
doc.add_paragraph('Implemented a hybrid predictive architecture that blends rule-based behavioral scoring (workload, sleep deficit) with ML predictions to generate a highly accurate Wellness Index.', style='List Bullet')

doc.add_heading('3.2 Aira Chatbot Assistant', level=3)
doc.add_paragraph('Built a conversational AI orchestrator equipped with contextual memory management to remember user inputs.', style='List Bullet')
doc.add_paragraph('Implemented prompt builders, response validators, and crisis handling logic to provide empathetic, non-medical emotional support and awareness insights safely.', style='List Bullet')

doc.add_heading('3.3 Quality Assurance & Automated Testing', level=3)
doc.add_paragraph('Developed a comprehensive testing suite utilizing pytest to ensure system reliability.', style='List Bullet')
doc.add_paragraph('Executed over 350 unit and integration tests across core models, routes, and services to validate data integrity.', style='List Bullet')

doc.add_heading('4. CURRENT SYSTEM STATUS (FINAL VIEW)', level=2)
doc.add_paragraph('The AIRA system has evolved from a preliminary skeleton into a cohesive multi-layer application. The current status of the project components is as follows:')
doc.add_paragraph('Completed: AI model integration, backend API completion, emotion detection pipeline, automated test suite, dashboard data visualization, and MongoDB schema design.', style='List Bullet')
doc.add_paragraph('In Progress / Remaining: Finalizing geolocation integrations and preparing for deployment.', style='List Bullet')

doc.add_heading('5. PLAN FOR REMAINING PHASE (FINAL 3 DAYS)', level=2)
doc.add_paragraph('With the internship concluding in 3 days, the team will focus exclusively on the following closure tasks:')
doc.add_paragraph('1. Doctor Geolocation Adding: Finalizing and integrating the geolocation features (Haversine formula) to accurately compute distances and match students with nearby mental health professionals.', style='List Bullet')
doc.add_paragraph('2. Testing: Conducting the final rounds of comprehensive system testing (UI, integration, and security checks), bug fixing, and quality assurance to ensure all modules run smoothly.', style='List Bullet')
doc.add_paragraph('3. Deployment: Finalizing production environment configurations (Gunicorn), resolving any dependency deprecations, and successfully deploying the AIRA platform to a live web hosting environment.', style='List Bullet')

doc.add_heading('6. CONCLUSION', level=2)
doc.add_paragraph('The AIRA project has successfully transitioned from an exploratory research concept into a functional AI-assisted mental wellness application. By fulfilling the objectives outlined in the milestone phase, the team has delivered a secure, responsive, and intelligent platform designed to provide actionable emotional awareness for students.')

doc.add_heading('Declaration', level=2)
doc.add_paragraph('We hereby declare that the work reported above reflects the collaborative effort of the project team and has been carried out in good faith as part of the AIRA project.')
doc.add_paragraph('\nSignatures: \n\n_______________________                  _______________________\nDiksha Shekhawat                         Anand Singh Rathore')

doc.save('Report_III_Project_Condition_Diksha_Shekhawat.docx')
print("Sequel team report document created successfully.")
