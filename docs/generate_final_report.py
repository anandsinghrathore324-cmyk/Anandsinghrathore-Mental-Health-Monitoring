import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = docx.Document()

    # Set Margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(6)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.paragraph_format.space_after = Pt(12)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(102, 102, 102)
        p.paragraph_format.space_after = Pt(24)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 153)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_p(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)
        return p

    # ---------------------------------------------------------
    # COVER PAGE / TITLE PAGE
    # ---------------------------------------------------------
    add_title("PRACTICE SCHOOL – I FINAL REPORT")
    add_subtitle("AIRA: AI-Based Student Mental Health Monitoring and Support System")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Submitted in partial fulfilment of the requirements for the degree of\nBachelor of Technology in Computer Science Engineering")
    r.font.size = Pt(12)
    r.font.italic = True
    p.paragraph_format.space_after = Pt(36)

    # Student details table
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PREPARED BY:")
    r.font.bold = True
    r.font.size = Pt(12)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Student Name"
    hdr[1].text = "Roll Number"
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True

    r1 = table.rows[1].cells
    r1[0].text = "Diksha Shekhawat"
    r1[1].text = "2024BTECH156"

    r2 = table.rows[2].cells
    r2[0].text = "Anand Singh Rathore"
    r2[1].text = "2024BTECH158"

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Supervisor details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FACULTY SUPERVISORS & EXAMINERS:\nDr. Sonali Vyas & Dr. Rajnish Kumar\n\nEXTERNAL SUPERVISOR:\nDr. Saurabh Kumar")
    r.font.size = Pt(11)
    r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Department of Computer Science Engineering\nInstitute of Engineering and Technology\nJK Lakshmipat University, Jaipur\nAugust 2026")
    r.font.size = Pt(11)

    doc.add_page_break()

    # ---------------------------------------------------------
    # CERTIFICATE OF WORK COMPLETION
    # ---------------------------------------------------------
    add_h1("CERTIFICATE OF WORK COMPLETION")
    add_p("This is to certify that the Practice School-I project report entitled \"AIRA: AI-Based Student Mental Health Monitoring and Support System\" submitted by Diksha Shekhawat (Roll No: 2024BTECH156) and Anand Singh Rathore (Roll No: 2024BTECH158) towards the partial fulfillment of the requirements for the degree of Bachelor of Technology in Computer Science Engineering of JK Lakshmipat University, Jaipur, is an authentic record of work carried out by them under our supervision and guidance.")
    add_p("In our opinion, the submitted work has reached the required academic and technical standard for being accepted for the Practice School-I examination.")

    doc.add_paragraph().paragraph_format.space_after = Pt(48)

    # Signatures
    p = doc.add_paragraph()
    p.add_run("_________________________\t\t_________________________\nDr. Sonali Vyas\t\t\t\tDr. Rajnish Kumar\nFaculty Supervisor\t\t\tDepartment of CSE, JKLU")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(36)

    p = doc.add_paragraph()
    p.add_run("_________________________\nDr. Saurabh Kumar\nExternal Supervisor")

    doc.add_page_break()

    # ---------------------------------------------------------
    # ACKNOWLEDGEMENTS
    # ---------------------------------------------------------
    add_h1("ACKNOWLEDGEMENTS")
    add_p("We express our profound gratitude to our supervisors and mentors who provided invaluable guidance, encouragement, and technical insight throughout the development of the AIRA Student Mental Health Monitoring Platform.")
    add_p("We are deeply grateful to Dr. Sonali Vyas and Dr. Rajnish Kumar from the Department of Computer Science Engineering at JK Lakshmipat University, Jaipur, for their continuous academic supervision, constructive feedback, and architectural evaluation.")
    add_p("We extend our sincere thanks to Dr. Saurabh Kumar, our external supervisor, for sharing industry standards, guiding our API integration workflows, and helping us align the system requirements with real-world software engineering practices.")
    add_p("Finally, we thank JK Lakshmipat University for providing the state-of-the-art software labs and infrastructure necessary to build, test, and deploy this project.")

    doc.add_page_break()

    # ---------------------------------------------------------
    # ABSTRACT / EXECUTIVE SUMMARY
    # ---------------------------------------------------------
    add_h1("ABSTRACT")
    add_p("Student mental health is a critical concern in modern higher education, where academic workload, exam stress, sleep deprivation, and emotional burnout severely impact student well-being. AIRA (AI-Based Student Mental Health Monitoring & Support System) is an intelligent full-stack web platform designed to detect early emotional distress, analyze behavioral patterns, provide real-time AI conversational support, and dynamically connect students to verified local mental health specialists.")
    add_p("AIRA integrates dual machine learning prediction engines: a natural language processing (NLP) model for analyzing free-text journal entries and a behavioral regression model for evaluating numerical stress factors. The system features a responsive cyberpunk-themed glassmorphism interface, an interactive 30-Day Mood Heatmap, an empathetic AI Chatbot Assistant with contextual memory, and a dynamic Geolocation Specialist Referral module powered by the Google Places API (New) and browser HTML5 GPS location auto-detection with spatial Haversine distance calculations.")
    add_p("The platform's reliability has been empirically verified through an automated test suite comprising 175 unit and integration tests (100% pass rate). AIRA represents a scalable, production-ready solution for proactive student wellness management.")

    doc.add_page_break()

    # ---------------------------------------------------------
    # TABLE OF CONTENTS
    # ---------------------------------------------------------
    add_h1("TABLE OF CONTENTS")
    add_bullet("1. Introduction & Project Vision .......................................................................... Page 5")
    add_bullet("2. Progress from Milestone Stage .................................................................... Page 6")
    add_bullet("3. System Architecture & Technical Implementation ..................................... Page 7")
    add_bullet("    3.1 Machine Learning Diagnostic Models (NLP & Behavioral Regression)")
    add_bullet("    3.2 AIRA Chatbot Orchestrator & Contextual Memory System")
    add_bullet("    3.3 Live Google Places API Geolocation & Spatial Haversine Distance Engine")
    add_bullet("    3.4 Database Architecture & MongoDB Indexing Schemas")
    add_bullet("    3.5 Security Policies & Brevo REST API Email Driver")
    add_bullet("4. Empirical Quality Assurance & Automated Test Suite (175 Tests) ........... Page 11")
    add_bullet("5. Key Learnings & Engineering Challenges Solved ........................................ Page 12")
    add_bullet("6. Conclusion & Future Scope ........................................................................... Page 13")
    add_bullet("7. References ........................................................................................................ Page 14")

    doc.add_page_break()

    # ---------------------------------------------------------
    # CHAPTER 1: INTRODUCTION
    # ---------------------------------------------------------
    add_h1("1. INTRODUCTION & PROJECT VISION")
    add_p("University and high-school students routinely experience acute stress, anxiety disorders, and depression driven by high academic expectations, competitive examinations, financial pressure, and lifestyle imbalances. Despite the prevalence of emotional distress, traditional institutional counseling services face low engagement due to social stigma, high appointment waiting times, and lack of immediate crisis assistance.")
    add_p("AIRA (AI Student Mental Health Monitoring and Support System) was engineered to address this gap by establishing a continuous, non-invasive, and accessible digital wellness environment. The platform offers students a private, judgment-free space to monitor their emotional states, receive instant AI diagnostics, engage with an empathetic chatbot assistant, and immediately locate verified local psychologists and psychiatrists.")

    add_h2("1.1 Project Objectives")
    add_bullet("Real-Time Sentiment Analysis: Process free-text journal logs to detect sentiment distributions (joy, sadness, anxiety, anger).")
    add_bullet("Behavioral Stress Diagnostics: Compute a calibrated Wellness Index based on study hours, sleep deficit, screen time, and academic load.")
    add_bullet("Empathetic AI Conversational Support: Provide round-the-clock supportive guidance while enforcing strict crisis handling protocols.")
    add_bullet("Live Spatial Specialist Referrals: Automatically locate real-world mental health professionals using live GPS auto-detection, Google Places API (New), and spatial Haversine distance sorting.")
    add_bullet("Production-Grade Reliability: Validate system integrity via a comprehensive 175-test automated pytest suite.")

    # ---------------------------------------------------------
    # CHAPTER 2: PROGRESS FROM MILESTONE STAGE
    # ---------------------------------------------------------
    add_h1("2. PROGRESS FROM MILESTONE STAGE")
    add_p("At the mid-term milestone evaluation, the AIRA project was a prototype featuring static UI mockups, partial database connectivity, and unintegrated ML models. Over the final implementation phase, the development team successfully transformed AIRA into a production-ready web platform.")

    t = doc.add_table(rows=5, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    hdr[0].text = "Module / Feature"
    hdr[1].text = "Milestone Prototype Status"
    hdr[2].text = "Final Phase Completed Status"
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ("AI Diagnostic Engine", "Unintegrated static risk scores", "Integrated DistilBERT NLP + Behavioral ML Regression"),
        ("AIRA Chatbot Assistant", "Basic keyword rule responses", "Contextual memory orchestrator with crisis detection"),
        ("Geolocation & Doctor Search", "Hardcoded static coordinate lists", "Live Google Places API + HTML5 GPS + Haversine distance engine"),
        ("Quality Assurance & Testing", "Manual ad-hoc testing", "175 automated unit and integration tests (100% pass rate)")
    ]

    for row_idx, (m, ms, fs) in enumerate(data, start=1):
        cells = t.rows[row_idx].cells
        cells[0].text = m
        cells[1].text = ms
        cells[2].text = fs

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ---------------------------------------------------------
    # CHAPTER 3: SYSTEM ARCHITECTURE & IMPLEMENTATION
    # ---------------------------------------------------------
    add_h1("3. SYSTEM ARCHITECTURE & TECHNICAL IMPLEMENTATION")
    add_p("AIRA is built on a decoupled full-stack architecture comprising a Vanilla ES6+ frontend client, a Flask Python 3.12 microservice backend, a MongoDB database layer, and external REST API integrations.")

    add_h2("3.1 Machine Learning Diagnostic Models")
    add_p("The diagnostic engine executes in-process via PredictionService, combining two core models:")
    add_bullet("Text Analysis Model (NLP): Utilizes a fine-tuned DistilBERT transformer pipeline to tokenize free-text logs and infer emotional sentiment probabilities across joy, sadness, anxiety, and neutral states.")
    add_bullet("Behavioral ML Regression Model: Processes quantitative lifestyle variables (sleep duration, study load, screen time, self-reported anxiety) to generate a calibrated Wellness Index (0-100) and risk tier (Low, Moderate, Elevated, Severe).")

    add_h2("3.2 AIRA Chatbot Orchestrator")
    add_p("The chatbot assistant is driven by ConversationOrchestrator, featuring:")
    add_bullet("Contextual Session Memory: Persists past dialogue turns per user ID to maintain coherent multi-turn conversations.")
    add_bullet("Crisis De-escalation Filter: Scans user inputs for self-harm or acute distress triggers, instantly displaying national crisis helpline numbers (e.g. Tele-MANAS 14416).")

    add_h2("3.3 Live Geolocation & Google Places API Engine")
    add_p("To eliminate static data limitations, AIRA incorporates a live spatial locator:")
    add_bullet("Browser HTML5 GPS Auto-Detection: Prompts for location permission on page load, capturing exact device coordinates and reverse-geocoding the city name.")
    add_bullet("Google Places API (New): Queries places.googleapis.com/v1/places:searchText for active mental health specialists using field mask optimization (displayName, rating, formattedAddress).")
    add_bullet("Haversine Distance Sorting: Computes exact spatial distances in kilometers from user coordinates. Strict fallback rules limit maximum distance to <= 100 km, preventing distant location leakage.")
    add_bullet("Dynamic Medical Avatars: Renders custom initial-based SVG avatars (ui-avatars.com) styled specifically for medical practitioners.")

    add_h2("3.4 Database Architecture & Security")
    add_bullet("MongoDB Collections: users, mental_health_reports, mood_logs, chatbot_history, doctor_recommendations, and otp_codes.")
    add_bullet("Brevo REST API Driver: Replaces blocked SMTP ports with HTTPS (port 443) REST dispatch for reliable 6-digit OTP delivery.")

    # ---------------------------------------------------------
    # CHAPTER 4: AUTOMATED TEST SUITE
    # ---------------------------------------------------------
    add_h1("4. EMPIRICAL QUALITY ASSURANCE & TEST SUITE")
    add_p("System stability was verified using an automated pytest suite (backend/tests/test_pytest_unit.py).")
    add_bullet("Total Executed Tests: 175 Tests")
    add_bullet("Pass Rate: 100% (175 Passed, 0 Failed)")
    add_bullet("Execution Time: ~20.67 Seconds")
    add_bullet("Test Coverage: Authentication endpoints, validation middleware, ML prediction pipelines, chatbot memory orchestrator, DoctorService Haversine calculations, Google Places API mocks, and MongoDB indexing constraints.")

    # ---------------------------------------------------------
    # CHAPTER 5: KEY LEARNINGS & CHALLENGES SOLVED
    # ---------------------------------------------------------
    add_h1("5. KEY LEARNINGS & ENGINEERING CHALLENGES SOLVED")
    add_h2("5.1 Overcoming External API Quota Limits (HTTP 429)")
    add_p("Challenge: During testing, Google Places API hit daily unbilled project quotas (HTTP 429 Resource Exhausted), causing fallback results to display unconstrained distant clinics.")
    add_p("Solution: Refactored DoctorService fallback logic to calculate Haversine distance for all fallback records and strictly cap results to <= 100 km, ensuring local relevance even during API downtime.")

    add_h2("5.2 Cloud SMTP Port Blocking")
    add_p("Challenge: Cloud hosting platforms (Render) block standard SMTP ports 25, 465, and 587.")
    add_p("Solution: Implemented a custom Brevo REST API email driver transmitting OTP verification payloads via HTTPS port 443.")

    add_h2("5.3 Key Technical Learnings")
    add_bullet("Designing modular Flask blueprints and RESTful endpoints.")
    add_bullet("Integrating transformer NLP and ML regression models into web API workflows.")
    add_bullet("Building responsive custom CSS glassmorphism UI without framework dependencies.")
    add_bullet("Implementing spatial algorithms (Haversine formula) and live API geocoding.")

    # ---------------------------------------------------------
    # CHAPTER 6: CONCLUSION & FUTURE SCOPE
    # ---------------------------------------------------------
    add_h1("6. CONCLUSION & FUTURE SCOPE")
    add_p("AIRA successfully demonstrates how artificial intelligence, behavioral analytics, and spatial web services can be combined to provide students with proactive, accessible, and life-saving mental health support. The platform is fully functional, thoroughly tested, and ready for deployment.")

    add_h2("6.1 Future Scope")
    add_bullet("Native Mobile Application: Developing React Native iOS and Android apps with push notifications for daily mood check-ins.")
    add_bullet("Wearable Sensor Integration: Syncing heart rate variability (HRV) data from smartwatches for physiological stress detection.")
    add_bullet("Tele-Counseling Video Appointments: Integrating WebRTC video call capabilities directly within the specialist referral module.")

    # ---------------------------------------------------------
    # REFERENCES
    # ---------------------------------------------------------
    add_h1("7. REFERENCES")
    refs = [
        "Aiven, 2024. Aiven Cloud Database Documentation. Available at: https://aiven.io/docs [Accessed 11 July 2026].",
        "Brevo, 2024. Brevo API v3 Documentation. Available at: https://developers.brevo.com/docs [Accessed 11 July 2026].",
        "Express.js / Flask Foundation, 2024. Flask Web Framework Documentation. Available at: https://flask.palletsprojects.com [Accessed 11 July 2026].",
        "Google Cloud Platform, 2024. Google Places API (New) Reference. Available at: https://developers.google.com/maps/documentation/places/web-service [Accessed 5 August 2026].",
        "Hugging Face, 2024. Transformers and DistilBERT Documentation. Available at: https://huggingface.co/docs/transformers [Accessed 11 July 2026].",
        "JSON Web Tokens, 2024. Introduction to JSON Web Tokens. Auth0 Inc. Available at: https://jwt.io/introduction [Accessed 11 July 2026].",
        "MongoDB Inc., 2024. PyMongo 4.0 Manual. Available at: https://pymongo.readthedocs.io [Accessed 11 July 2026].",
        "Pytest Development Team, 2024. Pytest Documentation. Available at: https://docs.pytest.org [Accessed 5 August 2026]."
    ]
    for r in refs:
        add_bullet(r)

    # Save docx
    docx_path = 'docs/Final_Report.docx'
    doc.save(docx_path)
    print(f"Successfully generated {docx_path}!")

    # Also generate Final_Report.md
    md_content = """# PRACTICE SCHOOL – I FINAL REPORT
## AIRA: AI-Based Student Mental Health Monitoring and Support System

**Degree**: Bachelor of Technology in Computer Science Engineering  
**Institution**: Department of Computer Science Engineering, Institute of Engineering and Technology, JK Lakshmipat University, Jaipur  
**Date**: August 2026  

---

### PREPARED BY:
- **Diksha Shekhawat** (Roll No: 2024BTECH156)
- **Anand Singh Rathore** (Roll No: 2024BTECH158)

### SUPERVISORS:
- **Faculty Supervisors**: Dr. Sonali Vyas & Dr. Rajnish Kumar
- **External Supervisor**: Dr. Saurabh Kumar

---

## CERTIFICATE OF WORK COMPLETION
This is to certify that the Practice School-I project report entitled **"AIRA: AI-Based Student Mental Health Monitoring and Support System"** submitted by **Diksha Shekhawat (2024BTECH156)** and **Anand Singh Rathore (2024BTECH158)** towards the partial fulfillment of the requirements for the degree of Bachelor of Technology in Computer Science Engineering of JK Lakshmipat University, Jaipur, is an authentic record of work carried out by them under our supervision and guidance.

In our opinion, the submitted work has reached the required academic and technical standard for being accepted for the Practice School-I examination.

---

## ACKNOWLEDGEMENTS
We express our profound gratitude to our supervisors Dr. Sonali Vyas, Dr. Rajnish Kumar, and Dr. Saurabh Kumar for their continuous academic supervision, constructive feedback, and architectural evaluation throughout the development of the AIRA platform.

---

## ABSTRACT
Student mental health is a critical concern in modern higher education. AIRA is an intelligent full-stack web platform designed to detect early emotional distress, analyze behavioral patterns, provide real-time AI conversational support, and dynamically connect students to verified local mental health specialists.

AIRA integrates dual machine learning prediction engines (DistilBERT NLP sentiment analysis + Behavioral ML regression), a responsive cyberpunk-themed interface, a 30-Day Mood Heatmap, an empathetic AI Chatbot Assistant, and a dynamic Geolocation Specialist Referral module powered by the Google Places API (New) and browser HTML5 GPS location auto-detection with spatial Haversine distance calculations.

System reliability is empirically verified through an automated test suite comprising 175 unit and integration tests (100% pass rate).

---

## TABLE OF CONTENTS
1. Introduction & Project Vision
2. Progress from Milestone Stage
3. System Architecture & Technical Implementation
   - 3.1 Machine Learning Diagnostic Models (NLP & Behavioral Regression)
   - 3.2 AIRA Chatbot Orchestrator & Contextual Memory System
   - 3.3 Live Google Places API Geolocation & Spatial Haversine Distance Engine
   - 3.4 Database Architecture & MongoDB Indexing Schemas
   - 3.5 Security Policies & Brevo REST API Email Driver
4. Empirical Quality Assurance & Automated Test Suite (175 Tests)
5. Key Learnings & Engineering Challenges Solved
6. Conclusion & Future Scope
7. References

---

## 1. INTRODUCTION & PROJECT VISION
University and high-school students routinely experience acute stress, anxiety disorders, and depression driven by academic load and lifestyle imbalances. AIRA provides a non-invasive, accessible digital wellness environment offering real-time AI diagnostics, empathetic chatbot support, and live specialist referrals.

---

## 2. PROGRESS FROM MILESTONE STAGE
| Module / Feature | Milestone Prototype Status | Final Phase Completed Status |
| :--- | :--- | :--- |
| **AI Diagnostic Engine** | Unintegrated static risk scores | Integrated DistilBERT NLP + Behavioral ML Regression |
| **AIRA Chatbot Assistant** | Basic keyword rule responses | Contextual memory orchestrator with crisis detection |
| **Geolocation & Doctor Search** | Hardcoded static coordinate lists | Live Google Places API + HTML5 GPS + Haversine distance engine |
| **Quality Assurance & Testing** | Manual ad-hoc testing | 175 automated unit and integration tests (100% pass rate) |

---

## 3. SYSTEM ARCHITECTURE & TECHNICAL IMPLEMENTATION
AIRA is built on a decoupled full-stack architecture:
- **Frontend**: Vanilla HTML5, CSS3 Glassmorphic UI, Vanilla ES6+ JS.
- **Backend**: Python 3.12, Flask Framework, PyJWT, Bcrypt, Brevo REST API.
- **Machine Learning**: Fine-tuned DistilBERT NLP transformer + Behavioral ML Regression model.
- **Geolocation**: Live Google Places API (New) (`places.googleapis.com/v1/places:searchText`), Browser HTML5 Geolocation (`navigator.geolocation`), Haversine distance calculation, dynamic SVG avatars (`ui-avatars.com`).
- **Database**: MongoDB with programmatic indexing on users, mental health reports, mood logs, OTP codes (TTL 5 mins), and doctor recommendations.

---

## 4. EMPIRICAL QUALITY ASSURANCE & AUTOMATED TEST SUITE
System stability is backed by an automated pytest suite:
- **Total Executed Tests**: 175 Tests
- **Pass Rate**: 100% (175 Passed, 0 Failed)
- **Execution Time**: ~20.67 Seconds

---

## 5. KEY LEARNINGS & ENGINEERING CHALLENGES SOLVED
- **Google Places API Quota Management (HTTP 429)**: Implemented Haversine distance calculations on fallback database records, strictly capping distance to `<= 100 km` to prevent distant location leakage.
- **Cloud SMTP Port Blocking**: Solved Render SMTP port restrictions by building a custom Brevo REST API driver over HTTPS (port 443).

---

## 6. CONCLUSION & FUTURE SCOPE
AIRA is a fully functional, empirically verified, and production-ready mental health monitoring platform. Future scope includes native mobile applications (React Native), smartwatch HRV wearable integration, and WebRTC video counseling.

---

## 7. REFERENCES
- Brevo API v3 Documentation: https://developers.brevo.com/docs
- Flask Web Framework: https://flask.palletsprojects.com
- Google Places API (New): https://developers.google.com/maps/documentation/places/web-service
- Hugging Face Transformers & DistilBERT: https://huggingface.co/docs/transformers
- Pytest Documentation: https://docs.pytest.org
"""
    md_path = 'docs/Final_Report.md'
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f"Successfully generated {md_path}!")

if __name__ == '__main__':
    create_report()
